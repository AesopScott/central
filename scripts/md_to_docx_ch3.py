"""Convert Chapter 3 Complete markdown to a styled .docx.

Handles: title heading, body paragraphs, <sub>Source...</sub> citation
lines (rendered small/gray), inline markdown links [text](url) -> real
Word hyperlinks, and *italic* runs.
"""
import re
import docx
from docx.shared import Pt, RGBColor
from docx.oxml.shared import OxmlElement, qn

SRC = r"g:\My Drive\central\dem\Chapter 3\10 Chapter 3 - Complete.md"
OUT = r"g:\My Drive\central\dem\Chapter 3\10 Chapter 3 - Complete.docx"

LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
TOKEN_RE = re.compile(r"(\[[^\]]+\]\([^)]+\)|\*[^*]+\*)")


def add_hyperlink(paragraph, url, text, small=False, color="0563C1"):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    if small:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "16")  # 8pt
        rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_runs(paragraph, text, small=False, gray=False):
    """Add text to paragraph, parsing links and *italics* into runs."""
    pos = 0
    for m in TOKEN_RE.finditer(text):
        if m.start() > pos:
            _plain(paragraph, text[pos:m.start()], small, gray)
        tok = m.group(0)
        link = LINK_RE.fullmatch(tok)
        if link:
            add_hyperlink(paragraph, link.group(2), link.group(1), small=small)
        elif tok.startswith("*"):
            r = paragraph.add_run(tok[1:-1])
            r.italic = True
            _style(r, small, gray)
        pos = m.end()
    if pos < len(text):
        _plain(paragraph, text[pos:], small, gray)


def _plain(paragraph, text, small, gray):
    r = paragraph.add_run(text)
    _style(r, small, gray)


def _style(run, small, gray):
    if small:
        run.font.size = Pt(8)
    if gray:
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def main():
    with open(SRC, encoding="utf-8") as f:
        lines = [ln.rstrip() for ln in f]

    doc = docx.Document()
    title_done = False
    for ln in lines:
        if not ln.strip():
            continue
        if not title_done:
            doc.add_heading(ln.strip(), level=0)
            title_done = True
            continue
        if ln.strip().startswith("<sub>"):
            body = ln.strip()
            body = body[len("<sub>"):]
            if body.endswith("</sub>"):
                body = body[:-len("</sub>")]
            p = doc.add_paragraph()
            add_runs(p, body, small=True, gray=True)
        else:
            p = doc.add_paragraph()
            add_runs(p, ln.strip())

    doc.save(OUT)
    print("WROTE", OUT)


if __name__ == "__main__":
    main()
