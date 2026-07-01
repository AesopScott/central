import re
from pathlib import Path

import docx
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.shared import qn
from docx.shared import Inches, Pt, RGBColor


CHAPTER_DIR = Path(r"G:\My Drive\central\dem\Chapter 4")
SRC = CHAPTER_DIR / "9 Chapter 4 - Council 3 Response.md"
MD_OUT = CHAPTER_DIR / "10 Chapter 4 - Complete.md"
DOCX_OUT = CHAPTER_DIR / "10 Chapter 4 - Complete.docx"

COMMENT_RE = re.compile(r"<!--.*?-->", re.DOTALL)
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
TOKEN_RE = re.compile(r"(\[\^\d+\]|\[[^\]]+\]\([^)]+\)|\*[^*]+\*)")


def clean_markdown(text: str) -> str:
    text = COMMENT_RE.sub("", text)
    lines = []
    for raw in text.splitlines():
        line = raw.rstrip()
        if line.strip() == "~~Feed my Worldview~~ The Self I Have Been Sold":
            line = "The Self I Have Been Sold"
        lines.append(line)

    cleaned = "\n".join(lines)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip() + "\n"
    return cleaned


def add_hyperlink(paragraph, url, text, *, small=False):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    if small:
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), "18")
        r_pr.append(size)
    run.append(r_pr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def style_run(run, *, small=False, gray=False):
    if small:
        run.font.size = Pt(9)
    if gray:
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_text_runs(paragraph, text, *, small=False, gray=False):
    pos = 0
    for match in TOKEN_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            style_run(run, small=small, gray=gray)
        token = match.group(0)
        link = LINK_RE.fullmatch(token)
        if link:
            add_hyperlink(paragraph, link.group(2), link.group(1), small=small)
        elif token.startswith("[^"):
            run = paragraph.add_run(token[2:-1])
            run.font.superscript = True
            style_run(run, small=small, gray=gray)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            style_run(run, small=small, gray=gray)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        style_run(run, small=small, gray=gray)


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Georgia"
    title.font.size = Pt(20)
    title.font.bold = False
    title.paragraph_format.space_after = Pt(18)


def build_docx(markdown: str):
    doc = docx.Document()
    configure_doc(doc)

    paragraphs = [p.strip() for p in markdown.split("\n\n") if p.strip()]
    title_done = False
    in_notes = False

    for block in paragraphs:
        if not title_done:
            para = doc.add_paragraph(style="Title")
            add_text_runs(para, block)
            title_done = True
            continue

        if block.startswith("[^"):
            if not in_notes:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                in_notes = True
            note = block.replace("\n", " ")
            note = re.sub(r"^\[\^(\d+)\]:\s*", r"[\1] ", note)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            add_text_runs(para, note, small=True, gray=True)
            continue

        if block.startswith("- "):
            for item in block.splitlines():
                para = doc.add_paragraph(style="List Bullet")
                add_text_runs(para, item[2:].strip())
            continue

        para = doc.add_paragraph()
        add_text_runs(para, block.replace("\n", " "))

    doc.save(DOCX_OUT)


def main():
    source = SRC.read_text(encoding="utf-8")
    cleaned = clean_markdown(source)
    MD_OUT.write_text(cleaned, encoding="utf-8")
    build_docx(cleaned)
    print(f"WROTE {MD_OUT}")
    print(f"WROTE {DOCX_OUT}")


if __name__ == "__main__":
    main()
